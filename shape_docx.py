from docx import Document
from docx.shared import Pt, Inches
from docx.oxml import OxmlElement
from docx.oxml.ns import qn

def set_cell_shading(cell, fill_color="D3D3D3"):  # light gray
    tcPr = cell._tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), fill_color)
    tcPr.append(shd)

def set_cell_borders(cell, color="555555", size=12):  # dark gray
    tcPr = cell._tc.get_or_add_tcPr()
    tblBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'left', 'bottom', 'right'):
        elem = OxmlElement(f'w:{edge}')
        elem.set(qn('w:val'), 'single')
        elem.set(qn('w:sz'), str(size))
        elem.set(qn('w:color'), color)
        tblBorders.append(elem)
    tcPr.append(tblBorders)

def create_shape_docx(filename):
    doc = Document()

    # 여백 약간 줄이기(선택)
    section = doc.sections[0]
    section.left_margin = Inches(0.7)
    section.right_margin = Inches(0.7)
    section.top_margin = Inches(0.7)
    section.bottom_margin = Inches(0.7)

    # 1x1 테이블로 '회색 박스' 만들기
    table = doc.add_table(rows=1, cols=1)
    table.autofit = False
    table.columns[0].width = Inches(6.5)  # 가로 폭 조정

    cell = table.cell(0, 0)
    set_cell_shading(cell, fill_color="D3D3D3")   # 배경색
    set_cell_borders(cell, color="555555", size=16)  # 테두리

    # 문단/텍스트
    p = cell.paragraphs[0]
    run = p.add_run("이 사각형 안에 들어있는 텍스트입니다.")
    run.font.name = "Malgun Gothic"       # 윈도우: 맑은 고딕
    run._element.rPr.rFonts.set(qn('w:eastAsia'), "Malgun Gothic")
    run.font.size = Pt(18)

    # 위아래 여백 비슷하게 보이도록 빈 줄 추가(옵션)
    cell.add_paragraph("")  # 아래쪽 여백용

    doc.save(filename)
    print(f"{filename} 생성 완료!")

if __name__ == "__main__":
    create_shape_docx("shape-demo.docx")
